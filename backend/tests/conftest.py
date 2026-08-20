import io
import os
import sys
from pathlib import Path

os.environ["DATABASE_URL"] = "sqlite:///./test_resume_screening.db"
os.environ.setdefault("GEMINI_API_KEY", "")

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest
from docx import Document
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas

from app.database.database import Base, engine
from app.database.init_db import init_db
from app.main import app


@pytest.fixture(scope="session", autouse=True)
def _setup_db():
    init_db()
    yield
    Base.metadata.drop_all(bind=engine)
    db_file = Path("test_resume_screening.db")
    if db_file.exists():
        try:
            db_file.unlink()
        except PermissionError:
            pass


@pytest.fixture()
def client():
    return TestClient(app)


def make_pdf_bytes(lines: list[str]) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    y = 800
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
    c.save()
    return buffer.getvalue()


def make_empty_pdf_bytes() -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    c.showPage()
    c.save()
    return buffer.getvalue()


def make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = Document()
    for para in paragraphs:
        document.add_paragraph(para)
    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row = table.add_row()
            for cell, value in zip(row.cells, row_values):
                cell.text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_empty_docx_bytes() -> bytes:
    document = Document()
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
