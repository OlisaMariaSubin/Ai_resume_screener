import io
import logging

from docx import Document

from app.utils.text_cleaning import normalize_text

logger = logging.getLogger(__name__)


def extract_text_from_docx(content: bytes) -> dict:
    """Extract text from a DOCX file's raw bytes, including table cell text
    (resumes frequently store skills/education/experience in tables).

    Returns {"success": True, "text": str} or {"success": False, "error": str}.
    """
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        logger.warning("Failed to open/parse DOCX: %s", exc)
        return {"success": False, "error": "Could not read DOCX file - it may be corrupted"}

    chunks: list[str] = []

    for para in document.paragraphs:
        if para.text.strip():
            chunks.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_texts:
                chunks.append(" | ".join(cell_texts))

    raw = "\n".join(chunks)
    cleaned = normalize_text(raw)

    if not cleaned:
        return {"success": False, "error": "No extractable text found in DOCX"}

    return {"success": True, "text": cleaned}
