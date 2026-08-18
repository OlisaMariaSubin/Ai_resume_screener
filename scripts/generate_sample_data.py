"""Generate synthetic sample resumes (DOCX) under data/sample_resumes/ and a
relevance-labelled evaluation set under data/evaluation/relevance.csv.

Relevance labels are derived directly from the designed skill overlap between
each resume and JD below (not hand-picked to inflate the eval score) - see
scripts/evaluate_precision.py for how they're consumed.
"""
import csv
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document  # noqa: E402

SAMPLE_RESUMES_DIR = ROOT / "data" / "sample_resumes"
SAMPLE_JDS_DIR = ROOT / "data" / "sample_jds"
EVALUATION_DIR = ROOT / "data" / "evaluation"

JD_REQUIRED_SKILLS = {
    "jd_backend_engineer": {"Python", "FastAPI", "SQL", "PostgreSQL", "Git", "Docker", "AWS", "Kubernetes"},
    "jd_frontend_engineer": {"JavaScript", "React", "HTML", "CSS", "TypeScript", "Next.js", "GraphQL"},
    "jd_data_scientist": {"Python", "Machine Learning", "Pandas", "NumPy", "SQL", "TensorFlow", "PyTorch", "scikit-learn"},
}

RESUMES = {
    "resume_backend_strong": {
        "name": "Priya Sharma",
        "email": "priya.sharma@example.com",
        "skills": ["Python", "FastAPI", "SQL", "PostgreSQL", "Git", "Docker", "AWS"],
        "education": ["Bachelor of Science in Computer Science, State University, 2026"],
        "experience": ["Backend Intern, Acme Corp - built REST APIs in Python (6 months)"],
    },
    "resume_backend_partial": {
        "name": "Diego Alvarez",
        "email": "diego.alvarez@example.com",
        "skills": ["Python", "SQL", "Git"],
        "education": ["Bachelor of Science in Information Technology, Tech College, 2026"],
        "experience": [],
    },
    "resume_backend_java": {
        "name": "Wei Zhang",
        "email": "wei.zhang@example.com",
        "skills": ["Java", "Spring Boot", "SQL", "Git", "Docker"],
        "education": ["Bachelor of Engineering, City University, 2025"],
        "experience": ["Software Intern, Beta Inc - Java backend services (3 months)"],
    },
    "resume_frontend_strong": {
        "name": "Sofia Ricci",
        "email": "sofia.ricci@example.com",
        "skills": ["JavaScript", "TypeScript", "React", "HTML", "CSS", "Next.js"],
        "education": ["Bachelor of Science in Computer Science, State University, 2026"],
        "experience": ["Frontend Intern, Gamma LLC - built React dashboards (4 months)"],
    },
    "resume_frontend_partial": {
        "name": "Omar Haddad",
        "email": "omar.haddad@example.com",
        "skills": ["JavaScript", "HTML", "CSS"],
        "education": ["Associate Degree in Web Development, Community College, 2025"],
        "experience": [],
    },
    "resume_frontend_backend_mix": {
        "name": "Emily Chen",
        "email": "emily.chen@example.com",
        "skills": ["JavaScript", "React", "Python", "SQL", "Git"],
        "education": ["Bachelor of Science in Computer Science, Metro University, 2026"],
        "experience": ["Full-stack Intern, Delta Co - React + Python (5 months)"],
    },
    "resume_data_strong": {
        "name": "Ananya Iyer",
        "email": "ananya.iyer@example.com",
        "skills": ["Python", "Machine Learning", "Pandas", "NumPy", "SQL", "scikit-learn"],
        "education": ["Master of Science in Data Science, Tech Institute, 2026"],
        "experience": ["Data Science Intern, Epsilon Analytics - ML pipelines (6 months)"],
    },
    "resume_data_partial": {
        "name": "Lucas Martin",
        "email": "lucas.martin@example.com",
        "skills": ["Python", "Pandas", "SQL"],
        "education": ["Bachelor of Science in Statistics, State University, 2026"],
        "experience": [],
    },
    "resume_data_deep_learning": {
        "name": "Hana Kobayashi",
        "email": "hana.kobayashi@example.com",
        "skills": ["Python", "Machine Learning", "TensorFlow", "PyTorch", "NumPy"],
        "education": ["Master of Science in Artificial Intelligence, Tech Institute, 2026"],
        "experience": ["Research Assistant - deep learning for vision (8 months)"],
    },
    "resume_unrelated": {
        "name": "Marcus Bell",
        "email": "marcus.bell@example.com",
        "skills": ["Project Management", "Agile", "Scrum"],
        "education": ["Bachelor of Business Administration, Coastal University, 2025"],
        "experience": ["Project Coordinator, Zeta Retail (1 year)"],
    },
}


def write_resume_docx(stem: str, data: dict) -> None:
    document = Document()
    document.add_paragraph(data["name"])
    document.add_paragraph(data["email"])

    document.add_paragraph("Skills")
    table = document.add_table(rows=0, cols=1)
    for skill in data["skills"]:
        row = table.add_row()
        row.cells[0].text = skill

    document.add_paragraph("Education")
    for edu in data["education"]:
        document.add_paragraph(edu)

    document.add_paragraph("Experience")
    for exp in data["experience"]:
        document.add_paragraph(exp)

    SAMPLE_RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    document.save(SAMPLE_RESUMES_DIR / f"{stem}.docx")


def relevance_label(resume_skills: set, jd_skills: set) -> int:
    """0-3 graded relevance, derived from actual skill overlap ratio."""
    if not jd_skills:
        return 0
    overlap_ratio = len(resume_skills & jd_skills) / len(jd_skills)
    if overlap_ratio >= 0.6:
        return 3
    if overlap_ratio >= 0.35:
        return 2
    if overlap_ratio > 0:
        return 1
    return 0


def main() -> None:
    for stem, data in RESUMES.items():
        write_resume_docx(stem, data)
    print(f"Wrote {len(RESUMES)} sample resumes to {SAMPLE_RESUMES_DIR}")

    EVALUATION_DIR.mkdir(parents=True, exist_ok=True)
    rows = []
    for jd_stem, jd_skills in JD_REQUIRED_SKILLS.items():
        for resume_stem, data in RESUMES.items():
            resume_skills = set(data["skills"])
            rows.append(
                {
                    "resume_id": resume_stem,
                    "jd_id": jd_stem,
                    "relevance": relevance_label(resume_skills, jd_skills),
                }
            )

    csv_path = EVALUATION_DIR / "relevance.csv"
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["resume_id", "jd_id", "relevance"])
        writer.writeheader()
        writer.writerows(rows)
    print(f"Wrote {len(rows)} labelled pairs to {csv_path}")


if __name__ == "__main__":
    main()
